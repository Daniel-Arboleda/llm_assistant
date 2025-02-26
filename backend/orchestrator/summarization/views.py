import os
import time
import logging
import torch
import re
from rest_framework.response import Response
from rest_framework.decorators import api_view
from django.conf import settings
from docx import Document
from .gpt_neo_loader import load_model, release_model

# Configuración de logs
logger = logging.getLogger('summarization')

# Directorios de almacenamiento
TRANSCRIPTS_DIR = os.path.abspath(os.path.join(settings.BASE_DIR, "downloads", "transcripts"))
SUMMARIES_DIR = os.path.abspath(os.path.join(settings.BASE_DIR, "downloads", "summaries"))

# Asegurar la existencia de las carpetas
os.makedirs(TRANSCRIPTS_DIR, exist_ok=True)
os.makedirs(SUMMARIES_DIR, exist_ok=True)

# Cargar el modelo y el tokenizador desde gpt_neo_loader.py
model, tokenizer = load_model()

def generate_summary(transcript_text):
    """Genera un resumen ejecutivo usando GPT-Neo."""
    try:
        start_time = time.time()
        prompt = ("Resumen ejecutivo de la siguiente reunión:\n\n" +
                  "---\n" + transcript_text[:1500] + "\n---\n" +
                  "Genera un resumen conciso en formato ejecutivo, resaltando los temas clave y conclusiones importantes.")
        
        inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(model.device)
        output = model.generate(
            **inputs,
            max_new_tokens=150,  
            min_length=80,   
            num_beams=5,      
            do_sample=False   
        )
        summary = tokenizer.decode(output[0], skip_special_tokens=True)
        elapsed_time = time.time() - start_time
        logger.info(f"Tiempo de generación del resumen: {elapsed_time:.4f} segundos")
        return summary
    except Exception as e:
        logger.error(f"Error al generar el resumen: {e}")
        return "No se pudo generar el resumen."

def save_summary_as_word(summary, file_name):
    """Guarda el resumen en un documento Word con formato APA."""
    doc = Document()
    doc.add_heading(f'Comité Directivo - Resumen Ejecutivo', level=1)
    doc.add_paragraph(f'Fecha: {time.strftime("%d/%m/%Y")}')
    doc.add_paragraph(f'Hora: {time.strftime("%H:%M")}')
    doc.add_paragraph('Lugar: Microsoft Teams')
    
    doc.add_heading('Asistencia de los miembros del Comité Directivo', level=2)
    doc.add_paragraph("Presidente - Asiste\nVicepresidente Financiero - Asiste\nVicepresidente Jurídico - Asiste")
    
    doc.add_heading('Temas tratados y conclusiones', level=2)
    doc.add_paragraph(summary)
    
    doc.add_heading('Puntos Clave de la Reunión', level=2)
    doc.add_paragraph("- Tema 1...\n- Tema 2...")
    
    doc.add_heading('Valores Importantes', level=2)
    doc.add_paragraph("- Métrica X: valor...\n- Métrica Y: valor...")
    
    doc_path = os.path.join(SUMMARIES_DIR, f"{file_name}.docx")
    doc.save(doc_path)
    return doc_path

@api_view(['POST'])
def process_transcript(request):
    """Recibe la transcripción, genera resumen y lo guarda en Word."""
    file_name = request.data.get("transcript_file")

    if not file_name or not isinstance(file_name, str) or ".." in file_name:
        logger.error("Nombre de archivo inválido o peligroso.")
        return Response({"error": "Nombre de archivo inválido."}, status=400)

    if not re.match(r'^[a-zA-Z0-9_\-.]+\.txt$', file_name):
        logger.error(f"Nombre de archivo inválido: {repr(file_name)}")
        return Response({"error": "Nombre de archivo inválido."}, status=400)

    transcript_path = os.path.join(TRANSCRIPTS_DIR, file_name)
    if not os.path.exists(transcript_path):
        logger.error(f"Archivo no encontrado: {transcript_path}")
        return Response({"error": "Archivo no encontrado."}, status=404)

    with open(transcript_path, "r", encoding="utf-8") as file:
        transcript_text = file.read()

    summary = generate_summary(transcript_text)
    doc_path = save_summary_as_word(summary, file_name.split('.')[0])

    return Response({
        "summary": summary,
        "word_document": doc_path
    }, status=200)
